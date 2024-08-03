from django.shortcuts import render
from django.http import HttpResponse
from transformers import GPT2LMHeadModel, GPT2Tokenizer
from huggingface_hub import login
from docx import Document
from io import BytesIO

# Log in to Hugging Face with your token
login(token="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")

# Load the model and tokenizer
model_name = "gpt2"  # Replace with your specific model if needed
model = GPT2LMHeadModel.from_pretrained(model_name)
tokenizer = GPT2Tokenizer.from_pretrained(model_name)

def generate_report_template(event_details):
    document = Document()
    document.add_heading('Event Report', level=1)
    document.add_paragraph(f"Event Name: {event_details['name']}")
    document.add_paragraph(f"Date: {event_details['date']}")
    document.add_paragraph(f"Venue: {event_details['venue']}")
    document.add_paragraph(f"Description:")
    document.add_paragraph(event_details['description'])
    document.add_paragraph(f"Attendees:")
    document.add_paragraph(event_details['attendees'])
    document.add_paragraph(f"Highlights:")
    document.add_paragraph(event_details['highlights'])
    return document

def generate_event_description(prompt):
    inputs = tokenizer.encode(prompt, return_tensors='pt')
    outputs = model.generate(
        inputs, 
        max_length=300, 
        num_return_sequences=1,
        no_repeat_ngram_size=2, 
        top_k=50, 
        top_p=0.95, 
        temperature=0.7
    )
    generated_text = tokenizer.decode(outputs[0], skip_special_tokens=True)
    return generated_text

def create_event_report(prompt):
    event_description = generate_event_description(prompt)
    event_details = {
        'name': 'Sample Event',
        'date': '2024-08-15',
        'venue': 'Sample Venue',
        'description': event_description,
        'attendees': 'Sample Attendees',
        'highlights': 'Sample Highlights'
    }
    document = generate_report_template(event_details)
    
    # Save the document to a BytesIO object
    doc_io = BytesIO()
    document.save(doc_io)
    doc_io.seek(0)
    return doc_io

def generate_report_view(request):
    if request.method == 'POST':
        prompt = request.POST.get('prompt')
        doc_io = create_event_report(prompt)
        
        response = HttpResponse(doc_io, content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document')
        response['Content-Disposition'] = 'attachment; filename=event_report.docx'
        return response
    return render(request, 'generate_report.html')
